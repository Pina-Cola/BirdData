#pip3 install python-docx

import psycopg2
from docx import Document
import glob



# connects to database
try:
    conn = psycopg2.connect("dbname='c5dv202_vt21_bio17sla' user='c5dv202_vt21_bio17sla' host='postgres.cs.umu.se' password='gf9uYTcrPusW'")
    print("I am connected to the database")
except:
    print("I am unable to connect to the database")


# now that connected to database, define a cursor to be able to execute commands
cur = conn.cursor()




def add_ringmarkning_2020(species, age, city, date_time, ring_id, place, file_path):

    try:
        cur.execute("INSERT INTO ringmarkning_2020_docx (ring_ID, species, age, place, city, date_time, file_path) VALUES (%s, %s, %s, %s, %s, %s, %s)",(ring_id, species, age, place, city, date_time, file_path))
        conn.commit() # <- We MUST commit to reflect the inserted data
    except:
        print("Could not execute command")

def add_ringmarkning_2021(species, age, city, date_time, ring_id, place, file_path):

    try:
        cur.execute("INSERT INTO ringmarkning_2021_docx (ring_ID, species, age, place, city, date_time, file_path) VALUES (%s, %s, %s, %s, %s, %s, %s)",(ring_id, species, age, place, city, date_time, file_path))
        conn.commit() # <- We MUST commit to reflect the inserted data
    except:
        print("Could not execute command")



ringmärkning_2020 = glob.glob("BirdData/Data/Ringmärkning/Kartor ringmärkning/Ringmärkning 2020/*")

for r in ringmärkning_2020:

    text = " "

    species = " "
    age = " "
    place = " "
    city = " "
    date_time = " "
    file_path = " "
    ring_id = " "

    document = Document(r)
    for p in document.paragraphs:
        text = p.text + text

    text = text.split(',')

    species = text[0]
    age = text[1][1:]
    place = text[2][1:]

    if "2020" not in text[4]:
        place = place + text[3]
        city = text[4][1:]
        text = text[5].split('\t')
    else:
        city = text[3][1:]
        text = text[4].split('\t')

    date_time = text[0]
    ring_id = text[1].split()[2]
    file_path = r

    file_path = " "

    add_ringmarkning_2020(species, age, city, date_time, ring_id, place, file_path)


    print("art=" + species + " ålder=" + age + " stad=" +  city + " datum=" + date_time + " ringID=" + ring_id  + " plats=" + place )


ringmärkning_2021 = glob.glob("BirdData/Data/Ringmärkning/Kartor ringmärkning/Ringmärkning 2021/*")

for r in ringmärkning_2021:

    text = " "

    species = " "
    age = " "
    place = " "
    city = " "
    date_time = " "
    file_path = " "
    ring_id = " "

    document = Document(r)
    for p in document.paragraphs:
        text = p.text + text

    text = text.split(',')

    species = text[0]
    age = text[1][1:]
    place = text[2][1:]

    if "2021" not in text[4]:
        place = place + text[3]
        city = text[4][1:]
        text = text[5].split('\t')
    else:
        city = text[3][1:]
        text = text[4].split('\t')

    date_time = text[0]
    ring_id = text[1].split()[2]
    file_path = r


    file_path = " "

    add_ringmarkning_2021(species, age, city, date_time, ring_id, place, file_path)

    print("art=" + species + " ålder=" + age + " stad=" +  city + " datum=" + date_time + " ringID=" + ring_id  + " plats=" + place )


# Link to maps
# http://www.google.com/maps/place/lat,lng
# http://www.google.com/maps/place/63.791861,19.904104
